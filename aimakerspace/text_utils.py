import os
from typing import List
import PyPDF2
from docx import Document


class TextFileLoader:
    def __init__(self, path: str, encoding: str = "utf-8"):
        self.documents = []
        self.path = path
        self.encoding = encoding

    def load(self):
        print(f"TextFileLoader.load() called for path: {self.path}")
        print(f"Path ends with .txt: {self.path.endswith('.txt')}")
        print(f"Is file: {os.path.isfile(self.path)}")
        print(f"Is directory: {os.path.isdir(self.path)}")
        
        if os.path.isdir(self.path):
            self.load_directory()
        elif os.path.isfile(self.path) and self.path.endswith(".txt"):
            self.load_file()
        else:
            raise ValueError(
                "Provided path is neither a valid directory nor a .txt file."
            )

    def load_file(self):
        # Try multiple encodings to handle problematic files
        encodings_to_try = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'latin-1', 'cp1252']
        
        for encoding in encodings_to_try:
            try:
                with open(self.path, "r", encoding=encoding) as f:
                    content = f.read()
                    # Clean up any problematic characters
                    content = content.encode('utf-8', errors='ignore').decode('utf-8')
                    self.documents.append(content)
                    return
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error reading file with {encoding} encoding: {e}")
                continue
        
        # If all encodings fail, try binary mode and decode with errors='ignore'
        try:
            with open(self.path, "rb") as f:
                content = f.read().decode('utf-8', errors='ignore')
                self.documents.append(content)
        except Exception as e:
            raise ValueError(f"Could not read file {self.path} with any encoding: {e}")

    def load_directory(self):
        for root, _, files in os.walk(self.path):
            for file in files:
                if file.endswith(".txt"):
                    try:
                        file_path = os.path.join(root, file)
                        # Try multiple encodings for each file
                        encodings_to_try = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'latin-1', 'cp1252']
                        
                        for encoding in encodings_to_try:
                            try:
                                with open(file_path, "r", encoding=encoding) as f:
                                    content = f.read()
                                    # Clean up any problematic characters
                                    content = content.encode('utf-8', errors='ignore').decode('utf-8')
                                    self.documents.append(content)
                                    break
                            except UnicodeDecodeError:
                                continue
                            except Exception as e:
                                print(f"Error reading {file_path} with {encoding} encoding: {e}")
                                continue
                        else:
                            # If all encodings fail, try binary mode
                            try:
                                with open(file_path, "rb") as f:
                                    content = f.read().decode('utf-8', errors='ignore')
                                    self.documents.append(content)
                            except Exception as e:
                                print(f"Could not read {file_path}: {e}")
                                continue
                    except Exception as e:
                        print(f"Error processing {file}: {e}")
                        continue

    def load_documents(self):
        self.load()
        return self.documents


class CharacterTextSplitter:
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        assert (
            chunk_size > chunk_overlap
        ), "Chunk size must be greater than chunk overlap"

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split(self, text: str) -> List[str]:
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunks.append(text[i : i + self.chunk_size])
        return chunks

    def split_texts(self, texts: List[str]) -> List[str]:
        chunks = []
        for text in texts:
            chunks.extend(self.split(text))
        return chunks


class PDFLoader:
    def __init__(self, path: str):
        self.documents = []
        self.path = path
        print(f"PDFLoader initialized with path: {self.path}")

    def load(self):
        print(f"Loading PDF from path: {self.path}")
        print(f"Path exists: {os.path.exists(self.path)}")
        print(f"Is file: {os.path.isfile(self.path)}")
        print(f"Is directory: {os.path.isdir(self.path)}")
        print(f"File permissions: {oct(os.stat(self.path).st_mode)[-3:]}")
        
        try:
            # Try to open the file first to verify access
            with open(self.path, 'rb') as test_file:
                pass
            
            # If we can open it, proceed with loading
            self.load_file()
            
        except IOError as e:
            raise ValueError(f"Cannot access file at '{self.path}': {str(e)}")
        except Exception as e:
            raise ValueError(f"Error processing file at '{self.path}': {str(e)}")

    def load_file(self):
        try:
            with open(self.path, 'rb') as file:
                # Create PDF reader object
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                text = ""
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Clean up any problematic characters
                            page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8')
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Error extracting text from PDF page: {e}")
                        continue
                
                if text.strip():
                    self.documents.append(text)
                else:
                    print(f"Warning: No text extracted from PDF {self.path}")
                    self.documents.append("")  # Add empty string to maintain structure
        except Exception as e:
            print(f"Error processing PDF {self.path}: {e}")
            # Add empty string to maintain structure even if PDF fails
            self.documents.append("")

    def load_directory(self):
        for root, _, files in os.walk(self.path):
            for file in files:
                if file.lower().endswith('.pdf'):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        
                        # Extract text from each page
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        
                        self.documents.append(text)

    def load_documents(self):
        self.load()
        return self.documents


class DOCXLoader:
    def __init__(self, path: str):
        self.documents = []
        self.path = path
        print(f"DOCXLoader initialized with path: {self.path}")

    def load(self):
        print(f"Loading DOCX from path: {self.path}")
        print(f"Path exists: {os.path.exists(self.path)}")
        print(f"Is file: {os.path.isfile(self.path)}")
        
        try:
            # Try to open the file first to verify access
            with open(self.path, 'rb') as test_file:
                pass
            
            # If we can open it, proceed with loading
            self.load_file()
            
        except IOError as e:
            raise ValueError(f"Cannot access file at '{self.path}': {str(e)}")
        except Exception as e:
            raise ValueError(f"Error processing file at '{self.path}': {str(e)}")

    def load_file(self):
        doc = Document(self.path)
        
        # Extract text from paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        self.documents.append(text)

    def load_documents(self):
        self.load()
        return self.documents


class MultiFileLoader:
    """Loader that can handle multiple file types (PDF, DOCX, TXT)"""
    
    def __init__(self):
        self.documents = []
        self.file_info = []  # Track which file each document came from
    
    def load_file(self, file_path: str, filename: str):
        """Load a single file based on its extension"""
        file_extension = filename.lower().split('.')[-1]
        print(f"Processing file: {filename} (extension: {file_extension}) at path: {file_path}")
        
        if file_extension == 'pdf':
            loader = PDFLoader(file_path)
        elif file_extension == 'docx':
            loader = DOCXLoader(file_path)
        elif file_extension == 'txt':
            loader = TextFileLoader(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Load the document
        docs = loader.load_documents()
        print(f"Loaded {len(docs)} documents from {filename}")
        
        # Add to our collection with file info
        for doc in docs:
            self.documents.append(doc)
            self.file_info.append(filename)
    
    def load_documents(self):
        return self.documents
    
    def get_file_info(self):
        """Return the list of filenames that were processed"""
        return self.file_info


if __name__ == "__main__":
    loader = TextFileLoader("data/KingLear.txt")
    loader.load()
    splitter = CharacterTextSplitter()
    chunks = splitter.split_texts(loader.documents)
    print(len(chunks))
    print(chunks[0])
    print("--------")
    print(chunks[1])
    print("--------")
    print(chunks[-2])
    print("--------")
    print(chunks[-1])
